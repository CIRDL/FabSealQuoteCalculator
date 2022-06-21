import quote
from quote import *
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_UNDERLINE
from datetime import date
from num2words import num2words


class DocumentCreator:
    def __init__(self, quote):
        self.quote = quote
        self.today_date = date.today().strftime("%B %d, %Y")
        self.doc = docx.Document()

    # Creates header section of document
    def create_header(self):
        # Create header paragraph for quote
        header = self.doc.add_paragraph(self.today_date)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Fill out important info
        header.add_run('\nNAME OF CONTACT')
        header.add_run('\nCOMPANY NAME')
        header.add_run('\nPHONE NUMBER')
        header.add_run('\nZip Code: FILL OUT HERE')
        header.add_run('\nCITY, STATE')

        # Double spacing for header
        header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Creates general body of document
    def create_general_body(self):
        # Records liner information
        body_string = num2words(self.quote.lining_system.total_liners).capitalize() + \
                      f"({self.quote.lining_system.total_liners}) liner"
        if self.quote.lining_system.total_liners > 1:
            body_string += "s"
        body_string += " fabricated from ENTER MATERIAL NAME HERE"
        body = self.doc.add_paragraph(body_string)
        self.__dimensions_info(body)
        self.__customizations_info(body)

    # Prints out dimensions based on tank shape
    def __dimensions_info(self, body):
        ref_liner = self.quote.lining_system.liner.info
        if isinstance(ref_liner, CLiner):
            body.add_run(f'\n\n{feet_deconverter(ref_liner.diameter_tank)}\''
                         f'-{inch_deconverter(ref_liner.diameter_tank)}\" diameter '
                         f'X {feet_deconverter(ref_liner.depth_tank)}\'-{inch_deconverter(ref_liner.depth_tank)}\" '
                         f'deep')
            if ref_liner.depth_liner - ref_liner.depth_tank > 0:
                body.add_run(" with ENTER DEPTH EXTENSIONS HERE. ")
            else:
                body.add_run(". ")

    # Prints out customizations to general body
    def __customizations_info(self, body):
        # Temp variable for shorter reference
        order_list = self.quote.accessories.orders
        # Iterates through order list if needed
        if len(order_list) > 0:
            body.add_run(f'Includes ')

            boots_doc = False
            sumps_doc = False
            manways_doc = False
            center_poles_doc = False
            columns_doc = False
            crates_doc = False

            # Loops through order list to add customizations to general body
            for order in order_list:

                # Sneaks in 'and'
                if len(order_list) > 1 and order_list[-2] == order:
                    body.add_run(f'and ')

                # Geo
                if isinstance(order, Geo):
                    if isinstance(self.quote.lining_system.liner.info, CLiner) or \
                            isinstance(self.quote.lining_system.liner.info, RLiner):
                        body.add_run(
                            f'16oz geotextile padding for the floor and {order.wall_thickness}oz '
                            f'geotextile padding for the sidewalls')
                        if order_list.index(order) != len(order_list) - 1:
                            body.add_run(", ")
                    # TODO - flat sheet option
                    else:
                        pass
                        # body.add_run(f"{num2words(geo_count)} ({geo_count}) layers of {material}oz geo")
                        # if order_list.index(order) != len(order_list) - 1:
                            # body.add_run(", ")

                # Batten strip
                elif isinstance(order, BattenStrips):
                    if order.type[0] == 'p':
                        body.add_run(f'poly-pro batten strips')
                        if order_list.index(order) != len(order_list) - 1:
                            body.add_run(", ")
                    else:
                        body.add_run(f'stainless steel batten strips')
                        if order_list.index(order) != len(order_list) - 1:
                            body.add_run(", ")

                # J-bolts
                elif isinstance(order, JBolts):
                    body.add_run(f'{num2words(order.jbolt_number)} ({order.jbolt_number}) j-bolts')
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Oarlocks
                elif isinstance(order, Oarlocks):
                    body.add_run(f'{num2words(order.oarlocks_number)} ({order.oarlocks_number}) oarlocks')
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Leak detection
                elif isinstance(order, LeakDetection):
                    body.add_run("leak detection")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Nailing strip
                elif isinstance(order, NailingStrip):
                    body.add_run('nailing strips')
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Stainless clips
                elif isinstance(order, StainlessClips):
                    body.add_run('stainless clips')
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # String check
                elif isinstance(order, str):
                    # Lifting hem
                    if order[0] == 'l' and order[1] == 'i':
                        body.add_run('lifting hem')
                        if order_list.index(order) != len(order_list) - 1:
                            body.add_run(", ")

                # Boots
                elif isinstance(order, Boot) and not boots_doc:
                    # Records documentation
                    boots_doc = True

                    # Create list of boot sizes
                    boot_list = self.quote.accessories.boot_list

                    # Create list of boot inches
                    boot_list_inches = []
                    for boot in boot_list:
                        boot_list_inches.append(boot.size)

                    # Create a set of the boot sizes
                    boot_set_inches = set(boot_list_inches)
                    # Turn it back into list for list properties with set abstracted goods
                    new_boot_list_inches = list(boot_set_inches)
                    new_boot_list_inches.sort()

                    # Iterate through set
                    for boot in new_boot_list_inches:
                        boot_count = 0
                        # Iterate through list
                        for x in boot_list_inches:
                            # Counts number of same type of boot
                            if x == boot:
                                boot_count += 1
                        # Sneaks in 'and' before last boot
                        if boot == boot_list_inches[-1] and order == order_list[-1] \
                                and (len(order_list) > 1 or len(new_boot_list_inches) > 1):
                            body.add_run(f'and ')
                        # Prints out info
                        body.add_run(f'{num2words(boot_count)} ({boot_count}) {boot}" boot')
                        if boot_count > 1:
                            body.add_run("s")
                        # Adds comma if more boots left, period otherwise
                        if boot == new_boot_list_inches[-1] and order == order_list[-1]:
                            body.add_run(".")
                        else:
                            body.add_run(", ")

                # Sumps
                elif isinstance(order, Sump) and not sumps_doc:
                    # Records documentation
                    sumps_doc = True

                    number_sumps = len(self.quote.accessories.sump_list)
                    body.add_run(f'{num2words(number_sumps)} ({number_sumps}) sump')
                    if number_sumps > 1:
                        body.add_run("s")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Manways
                elif isinstance(order, ManWay) and not manways_doc:
                    # Records documentation
                    manways_doc = True

                    number_manways = len(self.quote.accessories.manway_list)
                    body.add_run(f'{num2words(number_manways)} ({number_manways}) manway')
                    if number_manways > 1:
                        body.add_run("s")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Center poles
                elif isinstance(order, CenterPole) and not center_poles_doc:
                    # Records documentation
                    center_poles_doc = True

                    number_center_poles = len(self.quote.accessories.center_pole_list)
                    body.add_run(f'{num2words(number_center_poles)} ({number_center_poles}) center pole')
                    if number_center_poles > 1:
                        body.add_run("s")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Columns
                elif isinstance(order, Column) and not columns_doc:
                    # Records documentation
                    columns_doc = True

                    number_columns = len(self.quote.accessories.column_list)
                    body.add_run(f'{num2words(number_columns)} ({number_columns}) column')
                    if number_columns > 1:
                        body.add_run("s")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Crates
                elif isinstance(order, Crate) and not crates_doc:
                    # Records documentation
                    crates_doc = True

                    number_crates = len(self.quote.accessories.crate_list)
                    body.add_run(f'{num2words(number_crates)} ({number_crates}) crate')
                    if number_crates > 1:
                        body.add_run("s")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Installation
                elif isinstance(order, InstallationPackage):
                    body.add_run(f'installation')
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

                # Sneaks in '.'
                if order_list.index(order) == len(order_list) - 1 and order != 'boots':
                    body.add_run(".")

    # Creates calculations section of quote
    def create_calculations_body(self):
        ref_liner = self.quote.lining_system.liner.info
        if isinstance(ref_liner, CLiner):
            calculations_feet = self.doc.add_paragraph("\nBottom square footage:                         "
                                                    "                                                      "
                                                    "{:,}'".format(round(self.quote.lining_system.liner.info.
                                                                   liner_bottom_square_footage())))
            calculations_feet.add_run("\nSidewall square footage:                  "
                                      "                                                           ")
            square_footage_underline = calculations_feet.add_run("{:,}'".format(
                round(self.quote.lining_system.liner.info.liner_wall_square_footage())))
            # If lifting hem is ordered prints out square footage
            if self.quote.lining_system.liner.lifting_hem_area > 0:
                calculations_feet.add_run("\nLifting hem:                                              "
                                          "                                                         ")
                calculations_feet.add_run("{:,}'".format(round(self.quote.lining_system.liner.
                                                               lifting_hem_area))).underline = True

        # Prints out tank square footage
        calculations_feet.add_run("\nSquare footage:                                                "
                                  "                                              "
                                  "{:,}'".format(round(self.quote.lining_system.liner.total_liner_square_footage)))
        if not isinstance(ref_liner, FLiner):
            # If not flat sheet
            calculations_feet.add_run("\n5%:                                                         "
                                    "                                                            ")
            calculations_feet.add_run("{:,}'".format(round(self.
                                                           quote.lining_system.liner.total_five_percent))).\
                underline = True

        # Prints out the total square footage
        total_sqft_underline = calculations_feet.add_run("\nTotal square footage:                                 "
                                                         "                                                  "
                                                         "{:,}'".
                                                         format(round(self.quote.
                                                                      lining_system.liner.quote_square_footage)))
        total_sqft_underline.underline = WD_UNDERLINE.SINGLE

        # Prints out cost of material
        calculations_feet.add_run("\nCost of material:                                       "
                                  "                                                       "
                                  "${:,.2f}".format(self.quote.lining_system.liner.sq_price))

        # Double spaces the paragraph
        calculations_feet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

        # Extensions subsection of LINING SYSTEM

        # Prints out liner cost (not modified)
        extensions_cost = self.doc.add_paragraph("\nLiner cost:                                            "
                                              "                                                        "
                                                  "${:,.2f}".format(self.quote.lining_system.liner.total_cost))

        # Prints out discount if there is one
        if self.quote.lining_system.discount_percentage > 0:
            # Discount percentage
            extensions_cost.add_run("\nDiscount (" + str(self.quote.lining_system.discount_percentage) + "%):        "
                                                                                                    "              "
                                                                                       "                             "
                                                                                       "                          "
                                                                                       "             "
                                                                                       "${:,.2f}".format(
                self.quote.lining_system.liner.total_cost - self.quote.lining_system.liner_cost))
            # New cost of liner
            extensions_cost.add_run("\nNew cost of liner:                                            "
                                    "                                           ${:,.2f}".format(
                self.quote.lining_system.liner_cost))

        # Prints out total liner cost if added liners
        if self.quote.lining_system.total_liners > 1:
            extensions_cost.add_run(
                "\nTotal cost of (" + str(self.quote.lining_system.total_liners) + ") liners:            "
                                                                                          "                 "
                                                          "                                                   "
                                                          "${:,.2f}".format(self.quote.lining_system.cost))
        self.__customizations_calculations(extensions_cost)

        # Price of one lining system
        lining_system_underline = extensions_cost.add_run(
            "\nTotal cost for one (1) lining system:                                                  "
            "${:,.2f}".format(self.quote.accessories.cost + self.quote.lining_system.get_liner_cost()))

        # Calculate liner addition configuration
        if self.quote.accessories.additional_liners:
            total_lining_system_cost = self.quote.accessories.cost * self.quote.lining_system.total_liners
            lining_system_underline = extensions_cost.add_run("\nTotal cost for " + num2words
                (self.quote.lining_system.total_liners) + " ("
                                                              + str(self.quote.lining_system.total_liners) +
                                                              ") lining systems:                 "
                                "                              ${:,.2f}".format(total_lining_system_cost))
        lining_system_underline.underline = WD_UNDERLINE.SINGLE

        # Double space LINING SYSTEM subsection
        extensions_cost.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Creates customization calculation section of quote
    def __customizations_calculations(self, extensions_cost):
        ref_liner = self.quote.lining_system.liner.info
        order_list = self.quote.accessories.orders
        boots_doc = False
        sumps_doc = False
        manways_doc = False
        center_poles_doc = False
        columns_doc = False
        # Check to see if customizations were ordered
        if len(order_list) > 0:

            # Loop through order list to add to LINING SYSTEM
            for order in order_list:

                # Geo
                if isinstance(order, Geo):
                    if isinstance(ref_liner, FLiner):
                        extensions_cost.add_run("\n(" + str(self.quote.accessories.geo_layers) + ") "
                                                "Layers of " + str(order.wall_thickness)
                                                + "oz geotextile:                                                      "
                                                  "                      ${:,.2f}".format(order.cost))
                    else:
                        extensions_cost.add_run("\nGeotextile:                                                       "
                                                "                                             ${:,.2f}".format(order.
                                                                                                               cost))

                # Batten strip
                elif isinstance(order, BattenStrips):
                    extensions_cost.add_run(
                        "\nBatten strips (" + order.type.capitalize() + "):                      "
                                                                        "                                     "
                                                                        "              ${:,.2f}".format(order.cost))

                # J-bolts
                elif isinstance(order, JBolts):
                    extensions_cost.add_run("\n(" + str(order.jbolt_number) + ") J-bolts:                              "
                                                                     "                          "
                                                                     "                                        ${:,.2f}"
                                            .format(order.cost))

                # Oarlocks
                elif isinstance(order, Oarlocks):
                    extensions_cost.add_run("\n(" + str(order.oarlocks_number) + ") Oarlocks:                         "
                                                                          "                                           "
                                                                          "                             ${:,.2f}"
                                            .format(order.cost))

                # Nailing strip
                elif isinstance(order, NailingStrip):
                    extensions_cost.add_run(
                        "\nNailing strips:                                                         "
                        "                                        ${:,.2f}".format(order.cost))

                # Stainless clips
                elif isinstance(order, StainlessClips):
                    extensions_cost.add_run(
                        "\nStainless clips:                                                              "
                        "                                 ${:,.2f}".format(order.cost))

                # Boots
                elif isinstance(order, Boot) and not boots_doc:
                    boots_doc = True
                    total_boot_cost = 0
                    for boot in self.quote.accessories.boot_list:
                        total_boot_cost += boot.cost
                    extensions_cost.add_run(
                        "\n(" + str(len(self.quote.accessories.boot_list)) + ") Boots:                                "
                                                             "                                         "
                                                             "                                ${:,.2f}"
                        .format(total_boot_cost))

                # Sumps
                elif isinstance(order, Sump) and not sumps_doc:
                    sumps_doc = True
                    total_sump_cost = 0
                    for sump in self.quote.accessories.sump_list:
                        total_sump_cost += sump.cost
                    extensions_cost.add_run("\n(" + str(len(self.quote.accessories.sump_list)) + ") Sumps:             "
                                                                    "                                                  "
                                                                    "                                        ${:,.2f}"
                                            .format(total_sump_cost))

                # Manways
                elif isinstance(order, ManWay) and not manways_doc:
                    manways_doc = True
                    total_manway_cost = 0
                    for manway in self.quote.accessories.manway_list:
                        total_manway_cost += manway.cost
                    extensions_cost.add_run(
                        "\n(" + str(len(self.quote.accessories.manway_list)) + ") Manways:                             "
                                                      "                                      "
                                                      "                               ${:,.2f}".format(
                            total_manway_cost))

                # Center poles
                elif isinstance(order, CenterPole) and not center_poles_doc:
                    center_poles_doc = True
                    total_center_pole_cost = 0
                    for center_pole in self.quote.accessories.center_pole_list:
                        total_center_pole_cost += center_pole.cost
                    extensions_cost.add_run(
                        "\n(" + str(len(self.quote.accessories.center_pole_list)) + ") Center pole:                    "
                                                           "                                              "
                                                           "                           ${:,.2f}".format(
                            total_center_pole_cost))

                # Columns
                elif isinstance(order, Column) and not columns_doc:
                    columns_doc = True
                    total_column_cost = 0
                    for column in self.quote.accessories.column_list:
                        total_column_cost += column.cost
                    extensions_cost.add_run("\n(" + str(len(self.quote.accessories.column_list)) + ") Columns:         "
                                                                    "                                                  "
                                                                    "                                       ${:,.2f}"
                                            .format(total_column_cost))

                # Leak detection
                elif isinstance(order, LeakDetection):
                    extensions_cost.add_run(
                        "\nLeak detection:                                                       "
                        "                                     ${:,.2f}".format(order.cost))

    # Saves quote and terminates program
    def save(self):
        self.doc.save("Quote.docx")
