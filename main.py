# @author Cesar Ramirez
# @program Fabseal Quote Calculator


import math
from calculations import *
from accessories import *
from error import *
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_UNDERLINE
from datetime import date
from num2words import num2words


restart = False
while True:
    if restart:
        print("----------------------------------------\n")
        restart = False

    # Collect kind of tank
    tank = input("Enter type of tank (circular, rectangular, or flat sheet): ").lower()
    empty_response = empty_string(tank)
    while empty_response or (tank[0] != 'c' and tank[0] != 'r' and tank[0] != 'f' and tank != "restart"):
        tank = input("Please enter type of tank (circular, rectangular, or flat sheet): ").lower()
        empty_response = empty_string(tank)
    if tank == "restart":
        restart = True
        continue


    # Collect price per square foot
    sqft_price = input("Enter price per square foot: $")
    empty_response = empty_string(sqft_price)
    sqft_price = empty_literal(empty_response, "Please enter price per square foot: $", sqft_price)
    if sqft_price[0] == "r":
        restart = True
        continue
    sqft_price = float(sqft_price)


    # Collect weight per square foot
    weight_sqft = input("Enter weight per square foot: ")
    empty_response = empty_string(weight_sqft)
    weight_sqft = empty_literal(empty_response, "Please enter weight per square foot: ", weight_sqft)
    if weight_sqft[0] == "r":
        restart = True
        continue
    weight_sqft = float(weight_sqft)

    # For customizations
    circular = False
    rectangular = False
    lifting_hem = False
    additional_liner_cost = 0
    number_liners = 1

    while True:

        # Restart
        if restart:
            break

        # Circular tank configuration
        if tank[0] == "c":

            # Diameter configuration
            diameter_ft = input("Enter diameter of tank (ft): ")
            empty_response = empty_string(diameter_ft)
            diameter_ft = empty_literal(empty_response, "Please enter diameter of tank (ft): ", diameter_ft)
            if diameter_ft[0] == "r":
                restart = True
                continue
            diameter_ft = float(diameter_ft)

            diameter_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(diameter_inch)
            diameter_inch = empty_literal(empty_response, "Please enter any remaining inches: ", diameter_inch)
            if diameter_inch[0] == "r":
                restart = True
                continue
            diameter_inch = float(diameter_inch)

            diameter_tank = converter(diameter_ft, diameter_inch)
            diameter_liner = diameter_modifier(diameter_tank)

            # Depth configuration
            depth_ft = input("Enter depth of tank (ft): ")
            empty_response = empty_string(depth_ft)
            depth_ft = empty_literal(empty_response, "Please enter depth of tank (ft): ", depth_ft)
            if depth_ft[0] == "r":
                restart = True
                continue
            depth_ft = float(depth_ft)

            depth_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(depth_inch)
            depth_inch = empty_literal(empty_response, "Please enter any remaining inches: ", depth_inch)
            if depth_inch[0] == "r":
                restart = True
                continue
            depth_inch = float(depth_inch)

            depth_tank = converter(depth_ft, depth_inch)
            # Extra depth extensions
            depth_extensions = input("Please enter any extensions in inches to depth: ")
            empty_response = empty_string(depth_extensions)
            depth_extensions = empty_literal(empty_response, "Please enter any extensions in inches to depth: ", depth_extensions)
            if depth_extensions[0] == "r":
                restart = True
                continue
            depth_extensions = float(depth_extensions)

            depth_extensions = converter(0, depth_extensions)
            depth_liner = depth_tank + depth_extensions

            # Extensions
            circumference_tank = diameter_tank * math.pi
            circumference_liner = diameter_liner * math.pi
            circular = True

            # New line
            print("\n")

            # Calculate square footage
            actual_square_footage = circular_tank_sq_footage(diameter_liner, depth_liner)
            # Add the 5%
            five_percent = actual_square_footage * 0.05
            square_footage = round(actual_square_footage + five_percent)

            # Calculate weight of liner
            liner_weight = circular_weight(diameter_liner, depth_liner, weight_sqft)

            # Cost of liner
            liner_cost = square_footage * sqft_price

        # Rectangular tank configuration
        elif tank[0] == "r":

            # Length configuration
            length_ft = input("Enter length of tank (ft): ")
            empty_response = empty_string(length_ft)
            length_ft = empty_literal(empty_response, "Please enter length of tank (ft): ", length_ft)
            if length_ft[0] == "r":
                restart = True
                continue
            length_ft = float(length_ft)

            length_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(length_inch)
            length_inch = empty_literal(empty_response, "Please enter any remaining inches: ", length_inch)
            if length_inch[0] == "r":
                restart = True
                continue
            length_inch = float(length_inch)

            length_tank = converter(length_ft, length_inch)
            length_liner = length_modifier(length_tank)

            # Width configuration
            width_ft = input("Enter width of tank (ft): ")
            empty_response = empty_string(width_ft)
            width_ft = empty_literal(empty_response, "Please enter width of tank (ft): ", width_ft)
            if width_ft[0] == "r":
                restart = True
                continue
            width_ft = float(width_ft)

            width_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(width_inch)
            width_inch = empty_literal(empty_response, "Please enter any remaining inches: ", width_inch)
            if width_inch[0] == "r":
                restart = True
                continue
            width_inch = float(width_inch)

            width_tank = converter(width_ft, width_inch)
            width_liner = width_tank

            # Depth configuration
            depth_ft = input("Enter depth of tank (ft): ")
            empty_response = empty_string(depth_ft)
            depth_ft = empty_literal(empty_response, "Please enter depth of tank (ft): ", depth_ft)
            if depth_ft[0] == "r":
                restart = True
                continue
            depth_ft = float(depth_ft)

            depth_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(depth_inch)
            depth_inch = empty_literal(empty_response, "Please enter any remaining inches: ", depth_inch)
            if depth_inch[0] == "r":
                restart = True
                continue
            depth_inch = float(depth_inch)

            depth_tank = converter(depth_ft, depth_inch)

            # Extra depth dimensions
            depth_extensions = input("Please enter any extensions in inches to depth: ")
            empty_response = empty_string(depth_extensions)
            depth_extensions= empty_literal(empty_response, "Please enter any extensions in inches to depth: ", depth_extensions)
            if depth_extensions[0] == "r":
                restart = True
                continue
            depth_extensions = float(depth_extensions)

            depth_extensions = converter(0, depth_extensions)
            depth_liner = depth_tank + depth_extensions

            # Extensions
            perimeter_tank = (length_tank + width_tank) * 2
            perimeter_liner = (length_liner + width_liner) * 2
            rectangular = True

            # New line
            print("\n")

            # Calculate square footage
            actual_square_footage = rectangular_tank_sq_footage(length_liner, width_liner, depth_liner)
            # Add the 5%
            five_percent = actual_square_footage * 0.05
            square_footage = round(five_percent + actual_square_footage)

            # Calculate weight of liner
            liner_weight = rectangular_weight(length_liner, width_liner, depth_liner, weight_sqft)

            # Cost of liner
            liner_cost = square_footage * sqft_price

        # Flat sheet configuration
        elif tank[0] == "f":

            # Length configuration
            length_ft = input("Enter length (ft): ")
            empty_response = empty_string(length_ft)
            length_ft = empty_literal(empty_response, "Please enter length (ft): ", length_ft)
            if length_ft[0] == "r":
                restart = True
                continue
            length_ft = float(length_ft)

            length_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(length_inch)
            length_inch = empty_literal(empty_response, "Please enter any remaining inches: ", length_inch)
            if length_inch[0] == "r":
                restart = True
                continue
            length_inch = float(length_inch)

            length_liner = converter(length_ft, length_inch)

            # Width configuration
            width_ft = input("Enter width (ft): ")
            empty_response = empty_string(width_ft)
            width_ft = empty_literal(empty_response, "Please enter width (ft): ", width_ft)
            if width_ft[0] == "r":
                restart = True
                continue
            width_ft = float(width_ft)

            width_inch = input("Enter any remaining inches: ")
            empty_response = empty_string(width_inch)
            width_inch = empty_literal(empty_response, "Please enter any remaining inches: ", width_inch)
            if width_inch[0] == "r":
                restart = True
                continue
            width_inch = float(width_inch)

            width_liner = converter(width_ft, width_inch)

            # New line
            print("\n")

            # Calculate square footage
            square_footage = round(length_liner * width_liner)
            print("Square footage: {:,} ft.".format(square_footage))

            # Calculate weight of liner
            liner_weight = square_footage * weight_sqft

            # Cost of liner
            liner_cost = square_footage * sqft_price

        break
    if restart:
        continue
    break


# Print out cost and weight of single liner
print("Quote cost of liner: ${:,.2f}".format(liner_cost))
print("Weight of single liner: {:,} lbs".format(liner_weight))

# Setup for customization loop
total_quote_cost = liner_cost
total_mobilization_cost = 0
lining_system_cost = 0
total_weight = liner_weight
total_liners = 1
order_list = []
satisfied = False
discounted = False
installation_included = False
crate_included = False

# Prompt for customization loop
print("\n\nCustomize order below:\nType \'help\' for options\nType \'back\' for menu\n------------------------\n")

# Customization loop
while not satisfied:

    # Reset
    back_button = False

    # Get user command
    command = input("> ").lower()

    # To finish the order
    if command == 'finish':
        satisfied = True
        print("\nOrder Completed")
        continue

    # Dummy proof
    if command == '':
        print("")
        continue

    # Geo
    elif command[0] == 'g' and command[1] == 'e':

        geo_satisfied = False

        # Flat sheet calculations
        if not circular and not rectangular:

            geo_list = []

            while not geo_satisfied:

                geo_material_type = input("\nEnter thickness in ounces (16 or 8): ")
                if geo_material_type[0] == 'b' or geo_material_type[0] == 'B':
                    back_button = True
                    geo_satisfied = True
                    continue
                else:
                    geo_material_type = int(geo_material_type)

                while geo_material_type != 16 and geo_material_type != 8:
                    geo_material_type = int(input("Please enter 16 or 8: "))

                geo_layers = int(input(f"How many layers of {geo_material_type}oz geo would you like to add: "))

                for geo in range(geo_layers):
                    geo_list.append(geo_material_type)

                more_geo = input("Would you like to add more geo (yes/no)? ").lower()
                if more_geo[0] == 'n':
                    geo_satisfied = True

            # Back button
            if back_button:
                print("\n")
                continue

            total_geo_cost = 0
            total_geo_weight = 0

            for geo in geo_list:
                total_geo_cost += geo_cost_flatsheet(geo, length_liner, width_liner)
                total_geo_weight += geo_weight_flatsheet(geo_material_type, length_liner, width_liner)

            # Quote documentation
            lining_system_cost += total_geo_cost

            print("\nCost of geo added: ${:,.2f}".format(total_geo_cost))
            print("Weight of geo added: {:,} lbs.\n".format(total_geo_weight))

        # Circular calculations
        if circular:
            geo_material_type = input("\nEnter wall thickness in ounces (16 or 8): ")
            if geo_material_type[0] == 'b' or geo_material_type[0] == 'B':
                print("")
                continue
            else:
                geo_material_type = int(geo_material_type)
            total_geo_cost = geo_cost_circular(geo_material_type, wall_circular_tank_sq_footage
            (diameter_liner, depth_liner), floor_circular_tank_sq_footage(diameter_liner))
            total_geo_weight = geo_weight_circular(geo_material_type, wall_circular_tank_sq_footage(diameter_liner, depth_liner),
                                                   floor_circular_tank_sq_footage(diameter_liner))

            # Quote documentation
            lining_system_cost += total_geo_cost

            print("\nCost of geo added: ${:,.2f}".format(total_geo_cost))
            print("Weight of geo added: {:,} lbs.\n".format(total_geo_weight))

        # Rectangular calculations
        if rectangular:
            geo_material_type = input("\nEnter wall thickness in ounces (16 or 8): ")
            if geo_material_type[0] == 'b' or geo_material_type[0] == 'B':
                print("")
                continue
            else:
                geo_material_type = int(geo_material_type)
            total_geo_cost = geo_cost_rectangular(geo_material_type,
                                                  wall_rectangular_tank_sq_footage(length_liner, width_liner, depth_liner),
                                                  floor_rectangular_tank_sq_footage(length_liner, width_liner))
            total_geo_weight = geo_weight_rectangular(geo_material_type, length_liner, width_liner, depth_liner)

            # Quote documentation
            lining_system_cost += total_geo_cost

            print("\nCost of geo added: ${:,.2f}".format(total_geo_cost))
            print("Weight of geo added: {:,} lbs.\n".format(total_geo_weight))

        # Calculates total costs and weights
        total_quote_cost += total_geo_cost
        total_weight += total_geo_weight

        # Keep track of order
        order_list.append("geo")

        # Prints out important final info
        print("Total weight: {:,} lbs.".format(round(total_weight)))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Batten strip
    elif command[0] == 'b' and command[1] == 'a':

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no batten strip option for flat sheets.\n")
            continue

        # Retrieve batten strip type
        batten_strip_type = input("\nPoly-pro or stainless steel: ").lower()
        if batten_strip_type[0] == 'b':
            print("")
            continue

        # Turn type to cost
        if batten_strip_type[0] == 'p':
            batten_strip_cost = 10.0
            batten_strip_type = "poly-pro"
        else:
            batten_strip_cost = 33.30
            batten_strip_type = "stainless steel"

        # Figure out type of tank
        if circular:
            total_batten_cost = circumference_liner * batten_strip_cost
        else:
            total_batten_cost = perimeter_liner * batten_strip_cost

        # Calculate total quote cost
        total_quote_cost += total_batten_cost

        # Keep track of order
        order_list.append("batten strip")

        # Quote documentation
        lining_system_cost += total_batten_cost

        # Print out important info
        print("\nCost of batten strip added: ${:,.2f}".format(total_batten_cost))
        print("\nTotal quote cost: ${:,.2f}\n".format(total_quote_cost))

    # J-Bolt
    elif command[0] == 'j':

        # Cost per j-bolt
        jbolt_cost = 9.5

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no j-bolt option for flat sheets.\n")
            continue

        # Circular calculation for number of j-bolts
        if circular:
            jbolt_num = math.ceil(circumference_liner / 1.5)

        # Rectangular calculation for number of j-bolts
        else:
            jbolt_num = math.ceil(perimeter_liner / 1.5)

        # Calculates total cost of j-bolts
        total_jbolt_cost = jbolt_cost * jbolt_num

        # Calculates total quote cost
        total_quote_cost += total_jbolt_cost

        # Keep track of order
        order_list.append("jbolt")

        # Quote documentation
        lining_system_cost += total_jbolt_cost

        # Print out important information
        print("\nCost of j-bolts added: ${:,.2f}".format(total_jbolt_cost))
        print("\nTotal quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Oarlocks
    elif command[0] == 'o' and command[1] == 'a':

        # Cost per oarlock
        oarlock_price = 9.00

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no oarlock option for flat sheets.\n")
            continue

        # Collect number of oarlocks
        oarlock_number = input("\nHow many oarlocks would you like to add: ")
        if oarlock_number[0] == 'b' or oarlock_number[0] == 'B':
            print("")
            continue
        else:
            oarlock_number = int(oarlock_number)

        # Calculates cost of oarlocks
        total_oarlock_cost = oarlock_number * oarlock_price

        # Calculates total quote cost
        total_quote_cost += total_oarlock_cost

        # Keep track of order
        order_list.append("oarlock")

        # Quote documentation
        lining_system_cost += total_oarlock_cost

        # Print out important information
        print("\nCost of oarlocks added: ${:,.2f}".format(total_oarlock_cost))
        print("\nTotal quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Crates
    elif command[0] == 'c' and command[1] == 'r':

        # For while loop
        crate_satisfied = False
        big_crate_requested = False
        small_crate_requested = False
        total_crate_cost = 0
        total_crate_weight = 0
        back_button = False


        # Collect size of crate
        crate_size_satisfied = False
        while not crate_size_satisfied:
            crate_size = input("\nEnter crate size (big or small): ").lower()

            # Back
            if crate_size == 'back':
                print("")
                back_button = True
                crate_size_satisfied = True
                crate_satisfied = True
                continue

            # Turns crate size into cost per crate size
            if crate_size[0] == 'b' and crate_size[1] == 'i':
                crate_cost = 650
                crate_size_satisfied = True
            elif crate_size[0] == 's':
                crate_cost = 250
                crate_size_satisfied = True

        # If back button is pressed
        if back_button:
            continue

        # Collect number of crates
        number_crates = int(input("How many crates would you like: "))

        # Calculates weight of the crate based on size
        crate_weight = find_crate_weight(crate_size)

        # Find total crate info
        total_crate_cost += number_crates * crate_cost
        total_crate_weight += number_crates * crate_weight

        # Back button
        if back_button:
            continue

        print("\nCost of crate(s) added: ${:,.2f}".format(total_crate_cost))
        print("Weight of crate(s) added: {:,} lbs.\n".format(total_crate_weight))

        # Calculate total weight and cost
        total_weight += total_crate_weight
        total_quote_cost += total_crate_cost
        total_mobilization_cost += total_crate_cost

        # Keep track of order
        order_list.append("crate")
        crate_included = True

        # Prints out final info
        print("Total weight: {:,} lbs.".format(total_weight))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Leak detection
    elif command[0] == 'l' and command[1] == 'e':

        # Price per square foot of leak detection
        leak_detection_price = 10.00

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no leak detection option for flat sheets.\n")
            continue

        # Price calculation for circular
        if circular and not rectangular:
            total_leak_detection_cost = circumference_liner * leak_detection_price

        # Price calculation for rectangular
        else:
            total_leak_detection_cost = perimeter_liner * leak_detection_price

        # Calculate total cost
        total_quote_cost += total_leak_detection_cost

        # Keep track of order
        order_list.append("leak detection")

        # Quote documentation
        lining_system_cost += total_leak_detection_cost

        # Print out final info
        print("\nCost of leak detection added: ${:,.2f}".format(total_leak_detection_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Nailing strip
    elif command[0] == 'n' and command[1] == 'a':

        # Price per nailing strip
        nailing_strip_price = 1.00

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no nailing strip option for flat sheets.\n")
            continue

        # Price calculation for circular
        if circular and not rectangular:
            total_nailing_strip_cost = circumference_liner * nailing_strip_price

        # Price calculation for rectangular
        else:
            total_nailing_strip_cost = perimeter_liner * nailing_strip_price

        # Calculate total cost
        total_quote_cost += total_nailing_strip_cost

        # Keep track of order
        order_list.append("nailing strip")

        # Quote documentation
        lining_system_cost += total_nailing_strip_cost

        # Print out final info
        print("\nCost of nailing strip added: ${:,.2f}".format(total_nailing_strip_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Stainless clips
    elif command[0] == 's' and command[1] == 't':

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no stainless clips option for flat sheets.\n")
            continue

        # Price per stainless clip
        stainless_clip_price = 6.00

        # Price calculation for circular
        if circular and not rectangular:
            # Calculate proper circumference
            circumference_liner_rounded = math.ceil(circumference_liner)

            total_stainless_clips_cost = circumference_liner_rounded * stainless_clip_price

        # Price calculation for rectangular
        else:
            # Calculate proper perimeter
            perimeter_liner_rounded = math.ceil(perimeter_liner)
            total_stainless_clips_cost = perimeter_liner_rounded * stainless_clip_price

        # Calculate total cost
        total_quote_cost += total_stainless_clips_cost

        # Keep track of order
        order_list.append("stainless clips")

        # Quote documentation
        lining_system_cost += total_stainless_clips_cost

        # Print out final info
        print("\nCost of stainless clips added: ${:,.2f}".format(total_stainless_clips_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Installation
    elif command[0] == 'i' and command[1] == 'n':

        # Site survey
        # Determine whether the site survey is domestic
        is_inside_usa = input("\nIs the site survey in the US (yes/no)? ").lower()
        if is_inside_usa[0] == 'b':
            print("")
            continue
        # Site survey is national
        if is_inside_usa[0] == 'y':
            site_survey_distance_under = input("Is the site survey within 600 miles (yes/no)? ").lower()
            if site_survey_distance_under[0] == 'y':
                site_survey_cost = 1500
            else:
                site_survey_cost = 2500
        # Site survey is international
        else:
            site_survey_cost = float(input("Enter cost of site survey: $"))

        # Installation documentation
        installation_included = True

        # Traveling
        traveling_cost = float(input("Enter cost of traveling: $"))

        # Price of tools and hardware
        single_tools_and_hardware_cost = liner_cost * 0.1
        if number_liners > 1:
            total_tools_and_hardware_cost = single_tools_and_hardware_cost * number_liners
        else:
            total_tools_and_hardware_cost = single_tools_and_hardware_cost

        # Calculate based on tank type
        if circular and not rectangular:
            single_liner_installation_cost = calc_installation_price(diameter_liner)
            if number_liners > 1:
                total_liner_installation_cost = single_liner_installation_cost * number_liners
            else:
                total_liner_installation_cost = single_liner_installation_cost

            print("\nCost of standard installation: ${:,.2f}".format(total_liner_installation_cost))
            installation_request = input("Would you like to change cost (yes/no)? ").lower()
            if installation_request[0] == 'y':
                single_liner_installation_cost = float(input("Enter installation cost: $"))

        else:
            width_installation_cost = calc_installation_price(width_liner)
            length_installation_cost = calc_installation_price(length_liner)
            single_liner_installation_cost = max(width_installation_cost, length_installation_cost)
            if number_liners > 1:
                total_liner_installation_cost = single_liner_installation_cost * number_liners
            else:
                total_liner_installation_cost = single_liner_installation_cost

            print("\nCost of standard installation: ${:,.2f}".format(total_liner_installation_cost))
            installation_request = input("Would you like to change cost (yes/no)? ").lower()
            if installation_request[0] == 'y':
                single_liner_installation_cost = float(input("Enter installation cost: $"))

        # Calculate total installation cost
        total_installation_cost = site_survey_cost + traveling_cost \
                                  + total_liner_installation_cost + total_liner_installation_cost

        # Calculate total costs
        total_quote_cost += total_installation_cost
        total_mobilization_cost += total_installation_cost

        # Keep track of order
        order_list.append("installation")

        # Print out final info
        print("\nCost of site survey added: ${:,.2f}".format(site_survey_cost))
        print("Cost of traveling added: ${:,.2f}".format(traveling_cost))
        print("Cost of tools and hardware added: ${:,.2f}".format(total_tools_and_hardware_cost))
        print("Cost of installation added: ${:,.2f}".format(total_liner_installation_cost))
        print("\nCost of total installation package added: ${:,.2f}".format(total_installation_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Lifting hem
    elif command[0] == 'l' and command[1] == 'i':

        # Response to flat sheet attempt
        if not circular and not rectangular:
            print("\nThere is no lifting hem option for flat sheets.\n")
            continue

        # For quote document
        lifting_hem = True

        # Lifting hem cost for circular
        if circular and not rectangular:
            lifting_hem_cost = circumference_liner * sqft_price
            lifting_hem_weight = circumference_liner * weight_sqft

        # Lifting hem cost for rectangular
        else:
            lifting_hem_cost = perimeter_liner * sqft_price
            lifting_hem_weight = perimeter_liner * weight_sqft

        # Calculate total cost
        total_quote_cost += lifting_hem_cost

        # Calculate total weight
        total_weight += lifting_hem_weight

        # Keep track of order
        order_list.append("lifting hem")

        # Print out final info
        print("\nCost of lifting hem added: ${:,.2f}\n".format(lifting_hem_cost))

    # Boots
    elif command[0] == 'b' and command[1] == 'o':

        happy = False

        boot_list_inches = []

        while not happy:

            # Collect size in inches
            boot_size_inches = input("\nEnter boot size in inches: ")
            if boot_size_inches[0] == 'b' or boot_size_inches[0] == 'B':
                print("")
                back_button = True
                happy = True
                continue
            else:
                boot_size_inches = float(boot_size_inches)

            # Collect number of these boots
            boot_number = int(input(f"How many {boot_size_inches} inch boots would you like: "))

            # Add to boot list to record sizes of boots
            for x in range(boot_number):
                boot_list_inches.append(boot_size_inches)

            # See if more boots required
            more_boots = input("\nWould you like to add more boots (yes/no)? ").lower()
            if more_boots[0] == 'n':
                happy = True

        # Back button
        if back_button:
            continue

        # OCD
        boot_list_inches.sort()

        # For loop prep
        total_boot_cost = 0
        count = 0

        # Calculates total cost of boots
        for boot in boot_list_inches:

                boot_price = find_boot_price(boot)
                total_boot_cost += boot_price
                count += 1

        # Sums total boot cost and total quote cost
        total_quote_cost += total_boot_cost

        # Keep track of order
        order_list.append("boots")

        # Quote documentation
        lining_system_cost += total_boot_cost

        # Prints out important info
        print("\nPrice of boot(s) added: ${:,.2f}".format(total_boot_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Sumps
    elif command[0] == 's' and command[1] == 'u':

        # Price of sump labor
        sump_labor_cost = 250

        # Collect square footage of material
        sump_square_footage = input("\nEnter square footage of material used for sump: ")
        if sump_square_footage[0] == 'b' or sump_square_footage[0] == 'B':
            print("")
            continue
        else:
            sump_square_footage = float(sump_square_footage)

        # Collect number of sumps
        number_sumps = int(input("Enter number of sumps you wish to add: "))

        # Calculate total cost of sumps
        sump_cost = sump_square_footage * sqft_price
        total_sump_cost = (sump_labor_cost + sump_cost) * number_sumps

        # Calculate total cost
        total_quote_cost += total_sump_cost

        # Keep track of order
        order_list.append("sump")

        # Quote documentation
        lining_system_cost += total_sump_cost

        # Print out final info
        print("\nCost of sump(s) added: ${:,.2f}".format(total_sump_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Manways
    elif command[0] == 'm' and command[1] == 'a':

        # Price of manways
        manway_labor_cost = 250

        # Collect square footage of material
        manway_square_footage = input("\nEnter square footage of material used for manway: ")
        if manway_square_footage[0] == 'b' or manway_square_footage[0] == 'B':
            print("")
            continue
        else:
            manway_square_footage = float(manway_square_footage)

        # Collect number of manways
        number_manways = int(input("Enter number of manways you wish to add: "))

        # Calculate total cost of manways
        manway_cost = manway_square_footage * sqft_price
        total_manway_cost = (manway_labor_cost + manway_cost) * number_manways

        # Calculate total cost
        total_quote_cost += total_manway_cost

        # Keep track of order
        order_list.append("manway")

        # Quote documentation
        lining_system_cost += total_manway_cost

        # Print out final info
        print("\nCost of manway(s) added: ${:,.2f}".format(total_manway_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Center poles
    elif command[0] == 'c' and command[1] == 'e':

        # Price of center poles
        center_pole_labor_cost = 250

        # Collect square footage of material
        center_pole_square_footage = input("\nEnter square footage of material used for center pole: ")
        if center_pole_square_footage[0] == 'b' or center_pole_square_footage[0] == 'B':
            print("")
            continue
        else:
            center_pole_square_footage = float(center_pole_square_footage)

        # Collect number of center poles
        number_center_poles = int(input("Enter number of center poles you wish to add: "))

        # Calculate total cost of center poles
        center_pole_cost = center_pole_square_footage * sqft_price
        total_center_pole_cost = (center_pole_labor_cost + center_pole_cost) * number_center_poles

        # Calculate total cost
        total_quote_cost += total_center_pole_cost

        # Keep track of order
        order_list.append("center pole")

        # Quote documentation
        lining_system_cost += total_center_pole_cost

        # Print out final info
        print("\nCost of additional center pole(s) added: ${:,.2f}".format(total_center_pole_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Columns
    elif command[0] == 'c' and command[1] == 'o':

        # Price for column labor
        column_labor_cost = 250

        # Collect square footage of material
        column_square_footage = input("\nEnter square footage of material used for column: ")
        if column_square_footage[0] == 'b' or column_square_footage[0] == 'B':
            print("")
            continue
        else:
            column_square_footage = float(column_square_footage)

        # Collect number of columns
        number_columns = int(input("Enter number of columns you wish to add: "))

        # Calculate total cost of center poles
        column_cost = column_square_footage * sqft_price
        total_column_cost = (column_labor_cost + column_cost) * number_columns

        # Keep track of order
        total_quote_cost += total_column_cost

        # Keep track of order
        order_list.append("column")

        # Quote documentation
        lining_system_cost += total_column_cost

        # Print out final info
        print("\nCost of additional center pole(s) added: ${:,.2f}".format(total_column_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Discount the liner price
    elif command[0] == 'd' and command[1] == 'i':

        # Collect desired amount of discount
        discount_amount_percentage = input("\nEnter percentage you wish to discount liner price: ")
        if discount_amount_percentage[0] == 'b' or discount_amount_percentage[0] == 'B':
            print("")
            continue
        else:
            discount_amount_percentage = float(discount_amount_percentage)
        discount_amount_number = discount_amount_percentage / 100

        # Calculate discounted liner cost
        amount_discounted = discount_amount_number * liner_cost
        discounted_liner_cost = liner_cost - amount_discounted

        # For final info
        discounted = True

        # Subtract old liner(s) cost
        if number_liners > 1:
            original_liners_cost = total_liners * liner_cost
        else:
            original_liners_cost = liner_costc

        total_quote_cost -= original_liners_cost

        # Add new liner(s) cost
        if number_liners > 1:
            new_liners_cost = total_liners * discounted_liner_cost

        else:
            new_liners_cost = discounted_liner_cost

        total_quote_cost += new_liners_cost

        # Prints out new info
        print("\nCost of new liner with " + str(discount_amount_percentage) +
              "% discount: ${:,.2f}".format(discounted_liner_cost))
        if number_liners > 1:
            print("Total new liners cost: ${:,.2f}".format(new_liners_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Adds requested number of liners
    elif command[0] == 'a' and command[1] == 'd':

        # Collect number of liners
        number_liners = input("\nEnter number of liners you wish to add: ")
        if number_liners[0] == 'b' or number_liners[0] == 'B':
            print("")
            continue
        else:
            number_liners = int(number_liners)

        # Calculate total cost of liners
        additional_liner_cost = liner_cost * number_liners
        if discounted:
            additional_liner_cost = discounted_liner_cost * number_liners

        # Adds total liners of order
        total_liners += number_liners

        # Calculate total cost
        total_quote_cost += additional_liner_cost

        # Calculate total weight
        total_weight += number_liners * weight_sqft

        # Print out final info
        print("\nCost of additional liners added: ${:,.2f}".format(additional_liner_cost))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Help commands
    elif command == 'help':

        # For circular and rectangular tanks
        if circular or rectangular:
            print("\nCustomizations available:"
                  "\n--------------------------\nGeo\nBatten Strips\nJ-bolts\nOarlocks"
                  "\nCrate(s)\nLeak Detection\nNailing Strip\nStainless Clips\nLifting Hem\nInstallation"
                  "\nBoots\nSumps\nManways\nCenter poles\nColumns\nAdd liner(s)"
                  "\nDiscount liner\n\nTo complete order enter \'finish\'\n\n")

        # Flat sheet
        else:
            print("\nCustomizations available:"
                  "\n--------------------------\nGeo\nAdd liner(s)\nDiscount liner\n"
                  "\n\nTo complete order enter \'finish\'\n\n")

    # For user mistake
    else:
        print("\nPlease enter help if you need command list\n")


# DOCUMENT SECTION


# Today's date
today = date.today()
date_pretty = today.strftime("%B %d, %Y")

# Create document object
quote = docx.Document()

# HEADER for quote

# Create header paragraph for quote
header = quote.add_paragraph(date_pretty)
header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Fill out important info
header.add_run('\nNAME OF CONTACT')
header.add_run('\nCOMPANY NAME')
header.add_run('\nPHONE NUMBER')
header.add_run('\nZip Code: FILL OUT HERE')
header.add_run('\nCITY, STATE')

# Double spacing for header
header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# GENERAL BODY for quote

# Shows number of liners ordered
if total_liners > 1:
    body = quote.add_paragraph(num2words(total_liners).capitalize() + f'({total_liners}) liners fabricated from ENTER'
                                                                  f' MATERIAL NAME HERE')
else:
    body = quote.add_paragraph(num2words(total_liners).capitalize() + f'({total_liners}) liner fabricated from ENTER'
                                                                      f' MATERIAL NAME HERE')

# Prints out dimensions based on tank shape
# Circular
if circular and not rectangular:
    body.add_run(f'\n\n{diameter_ft}\'-{diameter_inch}\" diameter X {depth_ft}\'-{depth_inch}\" deep')
    if depth_extensions > 0:
        body.add_run(" with ENTER DEPTH EXTENSIONS HERE. ")
    else:
        body.add_run(". ")
# Rectangular
elif rectangular and not circular:
    body.add_run(f'\n\n{length_ft}\'-{length_inch}\" long X {width_ft}\'-{width_inch}\" wide X {depth_ft}\''
                         f'-{depth_inch}\" deep')
    if depth_extensions > 0:
        body.add_run(" with ENTER DEPTH EXTENSIONS HERE. ")
    else:
        body.add_run(". ")
# Flat sheet
else:
    body.add_run(f'\n\n{length_ft}\'-{length_inch}\" long X {width_ft}\'-{width_inch}\" wide. ')

# Prints out customizations to general body
if len(order_list) > 0:

    body.add_run(f'Includes ')

    # Loops through order list to add customizations to general body
    for order in order_list:

        # Sneaks in 'and'
        if order_list[-1] == order and order != 'boots' and len(order_list) > 1:
            body.add_run(f'and ')

        # Geo
        if order[0] == 'g' and order[1] == 'e':
            if circular or rectangular:
                body.add_run(f'16oz geotextile padding for the floor and {geo_material_type}oz geotextile padding for the '
                             f'sidewalls')
                if order_list.index(order) != len(order_list) - 1:
                    body.add_run(", ")
            else:
                geo_set = set(geo_list)
                for material in geo_set:
                    geo_count = 0
                    for geo in geo_list:
                        if geo == material:
                            geo_count += 1
                    body.add_run(f"{num2words(geo_count)} ({geo_count}) layers of {material}oz geo")
                    if order_list.index(order) != len(order_list) - 1:
                        body.add_run(", ")

        # Batten strip
        elif order[0] == 'b' and order[1] == 'a':
            if batten_strip_type[0] == 'p':
                body.add_run(f'poly-pro batten strips')
                if order_list.index(order) != len(order_list) - 1:
                    body.add_run(", ")
            else:
                body.add_run(f'stainless steel batten strips')
                if order_list.index(order) != len(order_list) - 1:
                    body.add_run(", ")

        # J-bolts
        elif order[0] == 'j' and order[1] == 'b':
            body.add_run(f'{num2words(jbolt_num)} ({jbolt_num}) j-bolts')
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Oarlocks
        elif order[0] == 'o' and order[1] == 'a':
            body.add_run(f'{num2words(oarlock_number)} ({oarlock_number}) oarlocks')
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Leak detection
        elif order[0] == 'l' and order[1] == 'e':
            body.add_run("leak detection")
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Nailing strip
        elif order[0] == 'n' and order[1] == 'a':
            body.add_run('nailing strips')
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Stainless clips
        elif order[0] == 's' and order[1] == 't':
            body.add_run('stainless clips')
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Lifting hem
        elif order[0] == 'l' and order[1] == 'i':
            body.add_run('lifting hem')
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Boots
        elif order[0] == 'b' and order[1] == 'o':

            # Create a set of the boot sizes
            boot_set_inches = set(boot_list_inches)
            # Turn it back into list for list properties with set abstracted goods
            new_boot_list_inches = list(boot_set_inches)
            new_boot_list_inches.sort()

            # Iterate through set
            for boot in new_boot_list_inches:
                boot_count = 0
                # Iterate through list
                for x in boot_list_inches:
                    # Counts number of same type of boot
                    if x == boot:
                        boot_count += 1
                # Sneaks in 'and' before last boot
                if boot == boot_list_inches[-1] and order == order_list[-1] \
                        and (len(order_list) > 1 or len(new_boot_list_inches) > 1):
                    body.add_run(f'and ')
                # Prints out info
                body.add_run(f'{num2words(boot_count)} ({boot_count}) {boot}" boot')
                if boot_count > 1:
                    body.add_run("s")
                # Adds comma if more boots left, period otherwise
                if boot == new_boot_list_inches[-1] and order == order_list[-1]:
                    body.add_run(".")
                else:
                    body.add_run(", ")

        # Sumps
        elif order[0] == 's' and order[1] == 'u':
            body.add_run(f'{num2words(number_sumps)} ({number_sumps}) sump')
            if number_sumps > 1:
                body.add_run("s")
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Manways
        elif order[0] == 'm' and order[1] == 'a':
            body.add_run(f'{num2words(number_manways)} ({number_manways}) manway')
            if number_manways > 1:
                body.add_run("s")
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Center poles
        elif order[0] == 'c' and order[1] == 'e':
            body.add_run(f'{num2words(number_center_poles)} ({number_center_poles}) center pole')
            if number_center_poles > 1:
                body.add_run("s")
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Columns
        elif order[0] == 'c' and order[1] == 'o':
            body.add_run(f'{num2words(number_columns)} ({number_columns}) column')
            if number_columns > 1:
                body.add_run("s")
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Crates
        elif order[0] == 'c' and order[1] == 'r':
            body.add_run(f'{num2words(number_crates)} ({number_crates}) crate')
            if number_crates > 1:
                body.add_run("s")
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Installation
        elif order[0] == 'i' and order[1] == 'n':
            body.add_run(f'installation')
            if order_list.index(order) != len(order_list) - 1:
                body.add_run(", ")

        # Sneaks in '.'
        if order_list.index(order) == len(order_list) - 1 and order != 'boots':
            body.add_run(".")


# LINING SYSTEM section of quote

# Liner subsection of LINING SYSTEM

# Circular feet number calculations
if circular and not rectangular:
    bottom_sqft = round(diameter_liner ** 2)
    sidewall_sqft = round(circumference_liner * depth_liner)
    calculations_feet = quote.add_paragraph("\nBottom square footage:                         "
                                            "                                                      "
                                            "{:,}'".format(bottom_sqft))
    calculations_feet.add_run("\nSidewall square footage:                  "
                              "                                                           ")
    square_footage_underline = calculations_feet.add_run("{:,}'".format(sidewall_sqft))

# Rectangular feet number calculations
elif rectangular and not circular:
    bottom_sqft = round(length_liner * width_liner)
    sidewall_sqft = round((length_liner + width_liner) * 2 * depth_liner)
    calculations_feet = quote.add_paragraph("\nBottom square footage:                         "
                                            "                                                      "
                                            "{:,}'".format(bottom_sqft))
    calculations_feet.add_run("\nSidewall square footage:                  "
                              "                                                           ")
    square_footage_underline = calculations_feet.add_run("{:,}'".format(sidewall_sqft))

# Flat sheet number calculations
else:
    bottom_sqft = round(length_liner * width_liner)
    calculations_feet = quote.add_paragraph("\nSquare footage:                                        "
                                            "                                                     ")
    square_footage_underline = calculations_feet.add_run("{:,}'".format(bottom_sqft))

# If lifting hem is ordered prints out square footage
if lifting_hem:
    # Circular calculations
    if circular and not rectangular:
        lifting_hem_sqft = circumference_liner
        actual_square_footage += lifting_hem_sqft
    else:
        lifting_hem_sqft = perimeter_liner
        actual_square_footage += lifting_hem_sqft
    five_percent = actual_square_footage * 0.05
    square_footage = round(five_percent + actual_square_footage)

    calculations_feet.add_run("\nLifting hem:                                              "
                              "                                                         ")
    calculations_feet.add_run("{:,}'".format(round(lifting_hem_sqft))).underline = True
else:
    square_footage_underline.underline = True

# Prints out square footage before 5% (actual)
if circular or rectangular:

    calculations_feet.add_run("\nSquare footage:                                                "
                              "                                              "
                              "{:,}'".format(round(actual_square_footage)))

# Prints out 5% if not a flat sheet
if circular or rectangular:
    calculations_feet.add_run("\n5%:                                                         "
                              "                                                            ")
    calculations_feet.add_run("{:,}'".format(round(five_percent))).underline = True

# Prints out the total square footage
total_sqft_underline = calculations_feet.add_run("\nTotal square footage:                                 "
                          "                                                  "
                          "{:,}'".format(square_footage))
total_sqft_underline.underline = WD_UNDERLINE.SINGLE

# Prints out cost of material
calculations_feet.add_run("\nCost of material:                                       "
                          "                                                       "
                          "${:,.2f}".format(sqft_price))

# Double spaces the paragraph
calculations_feet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


# Extensions subsection of LINING SYSTEM

# Prints out liner cost (not modified)
single_liner_cost = square_footage * sqft_price
extensions_cost = quote.add_paragraph("\nLiner cost:                                            "
                                        "                                                        "
                                        "${:,.2f}".format(single_liner_cost))

# Prints out discount if there is one
if discounted:
    # Discount percentage
    extensions_cost.add_run("\nDiscount (" + str(discount_amount_percentage) + "%):                      "
                                                                            "                             "
                                                                            "                          "
                                                                            "             "
                                                                            "${:,.2f}".format(amount_discounted))
    # New cost of liner
    extensions_cost.add_run("\nNew cost of liner:                                            "
                              "                                           ${:,.2f}".format(discounted_liner_cost))

# Prints out total liner cost if added liners
if total_liners > 1:

    # Prints total discounted liner cost
    if discounted:
        discounted_total_liner_cost = total_liners * discounted_liner_cost
        extensions_cost.add_run("\nTotal cost of (" + str(total_liners) + ") liners:                             "
                                "                                               "
                                "${:,.2f}".format(discounted_liner_cost * total_liners))
    # Prints total liner cost
    else:
        extensions_cost.add_run("\nTotal cost of (" + str(total_liners) + ") liners:                            "
                                "                                               "
                                    "${:,.2f}".format(single_liner_cost * total_liners))

# Check to see if customizations were ordered
if len(order_list) > 0:

    # Loop through order list to add to LINING SYSTEM
    for order in order_list:

        # Geo
        if order[0] == 'g' and order[1] == 'e':
            if not circular and not rectangular:
                extensions_cost.add_run("\n(" + str(len(geo_list)) + ") Layers of " + str(geo_material_type)
                                        + "oz geotextile:                                                       "
                                        "                     ${:,.2f}".format(total_geo_cost))
            else:
                extensions_cost.add_run("\nGeotextile:                                                       "
                                    "                                             ${:,.2f}".format(total_geo_cost))

        # Batten strips
        elif order[0] == 'b' and order[1] == 'a':
            extensions_cost.add_run("\nBatten strips (" + batten_strip_type.capitalize() + "):                      "                                                                             
                                                                "                                                     "
                                                                              "${:,.2f}".format(total_batten_cost))

        # J-bolts
        elif order[0] == 'j' and order[1] == 'b':
            extensions_cost.add_run("\n(" + str(jbolt_num) + ") J-bolts:                                       "
                                                                                    "                       "
                                        "                                          ${:,.2f}".format(total_jbolt_cost))

        # Oarlocks
        elif order[0] == 'o' and order[1] == 'a':
            extensions_cost.add_run("\n(" + str(oarlock_number) + ") Oarlocks:                              "
                                                                  "                                       "
                                            "                            ${:,.2f}".format(total_oarlock_cost))

        # Nailing strips
        elif order[0] == 'n' and order[1] == 'a':
            extensions_cost.add_run("\nNailing strips:                                                               "
                                    "                                      ${:,.2f}".format(total_nailing_strip_cost))

        # Stainless clips
        elif order[0] == 's' and order[1] == 't':
            extensions_cost.add_run("\nStainless clips:                                                              "
                                    "                                    ${:,.2f}".format(total_stainless_clips_cost))

        # Boots
        elif order[0] == 'b' and order[1] == 'o':
            extensions_cost.add_run("\n(" + str(len(boot_list_inches)) + ") Boots:                                  "
                                                                         "                                       "
                                    "                                ${:,.2f}".format(total_boot_cost))

        # Sumps
        elif order[0] == 's' and order[1] == 'u':
            extensions_cost.add_run("\n(" + str(number_sumps) + ") Sumps:                                    "
                                                                "                                    "
                                    "                               ${:,.2f}".format(total_sump_cost))

        # Manways
        elif order[0] == 'm' and order[1] == 'a':
            extensions_cost.add_run("\n(" + str(number_manways) + ") Manways:                                   "
                                                                  "                                "
                                    "                               ${:,.2f}".format(total_manway_cost))

        # Center poles
        elif order[0] == 'c' and order[1] == 'e':
            extensions_cost.add_run("\n(" + str(number_center_poles) + ") Center pole:                             "
                                                                       "                                     "
                                    "                           ${:,.2f}".format(total_center_pole_cost))

        # Columns
        elif order[0] == 'c' and order[1] == 'o':
            extensions_cost.add_run("\n(" + str(number_columns) + ") Columns:                                  "
                                                                  "                                   "
                                    "                             ${:,.2f}".format(total_column_cost))

        # Leak detection
        elif order[0] == 'l' and order[1] == 'e':
            extensions_cost.add_run("\nLeak detection:                                                            "
                                    "                                     ${:,.2f}".format(total_leak_detection_cost))

# Price of one lining system
if discounted:
    lining_system_cost += (new_liners_cost / number_liners)
else:
    lining_system_cost += single_liner_cost
lining_system_underline = extensions_cost.add_run("\nTotal cost for one (1) lining system:                                                  "
                        "${:,.2f}".format(lining_system_cost))

# Calculate liner addition configuration
if total_liners > 1:
    total_lining_system_cost = lining_system_cost * total_liners
    lining_system_underline = extensions_cost.add_run("\nTotal cost for " + num2words(total_liners) + " ("
                            + str(total_liners) + ") lining systems:                             "
                            "                  ${:,.2f}".format(total_lining_system_cost))
    lining_system_underline.underline = WD_UNDERLINE.SINGLE
else:
    lining_system_underline.underline = WD_UNDERLINE.SINGLE

# Double space LINING SYSTEM subsection
extensions_cost.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


crate_recorded = False

# INSTALLATION section of quote

# See if necessary to include
if installation_included:

    # Prints out header for installation section
    installation_header_paragraph = quote.add_paragraph()
    installation_header_underline = installation_header_paragraph.add_run("\nInstallation for "
                                                    + num2words(total_liners) + f" ({total_liners}) lining systems:")
    # Paragraph formatting
    installation_header_underline.underline = True

    # Prints out installation cost
    installation_paragraph = quote.add_paragraph("Installation:                                                    "
                                   "                                            ${:,.2f}"
                                   .format(total_liner_installation_cost))

    # Prints out site survey
    installation_paragraph.add_run("\nSite survey:                                                           "
                                   "                                       ${:,.2f}".format(site_survey_cost))

    # Prints out travel costs
    installation_paragraph.add_run("\nAirfare/lodging/transportation:                                         "
                                   "                  ${:,.2f}".format(traveling_cost))

    # If crate is included
    # NOTE - should be last
    if crate_included:
        crate_recorded = True
        installation_paragraph.add_run("\n" + num2words(number_crates).capitalize() + " (" + str(number_crates) + ") crate")
        if number_crates > 1:
            installation_paragraph.add_run("s")
        installation_paragraph.add_run(":                                                               "
                                        "                              ${:,.2f}".format(total_crate_cost))

    # Report total mobilization cost
    installation_paragraph.add_run("\nTotal Mobilization Cost:                                                        "
                                   "                 ${:,.2f}".format(total_mobilization_cost))

    # Paragraph formatting
    installation_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


# INSTALLATION wrap up section of quote
# Create new paragraph for potential crate and turnkey cost
installation_wrap_up_paragraph = quote.add_paragraph()

# If crate(s) included but not installation
if crate_included and not crate_recorded:
    installation_wrap_up_paragraph.add_run("\n" + num2words(number_crates).capitalize()
                                            + " (" + str(number_crates) + ") crate")
    if number_crates > 1:
        installation_wrap_up_paragraph.add_run("s")
    installation_wrap_up_paragraph.add_run(":                                                               "
                                    "                                  ${:,.2f}".format(total_crate_cost))

# Calculates total turnkey cost
total_turnkey_cost = 0
if total_liners > 1:
    total_turnkey_cost += total_lining_system_cost
else:
    total_turnkey_cost += lining_system_cost
if installation_included:
    total_turnkey_cost += total_mobilization_cost
if crate_included and not crate_recorded:
    total_turnkey_cost += total_crate_cost

# Prints out total turnkey cost
turnkey_underline = installation_wrap_up_paragraph.add_run("\n\nTotal Turnkey Cost:                           "
                     "                                                      ${:,.2f}".format(total_turnkey_cost))
turnkey_underline.underline = WD_UNDERLINE.SINGLE

# Paragraph formatting
installation_wrap_up_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# WEIGTH section of quote
weight_section = quote.add_paragraph("Total Weight:                                                 "
                                    "                                              {:,}lbs".format(round(total_weight)))

# Save quote document as quote
quote.save("Quote.docx")
